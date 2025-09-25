from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement

# Create a new Word document
doc = Document()

# Define style helpers
def set_run_style(run, bold=False, font_size=11, color=(0, 0, 0)):
    run.font.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    run.font.color.rgb = RGBColor(*color)

def add_heading(text, size=14, color=(0, 102, 204)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_style(run, bold=True, font_size=size, color=color)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def add_section_title(title):
    doc.add_paragraph()
    add_heading(f"● {title.upper()}", size=12)

def add_bullet_point(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)

def add_paragraph(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_style(run, bold=bold, font_size=11)

# Header
add_heading("ELHUSSEINI ABDELFATTAH", size=16)
add_paragraph("Licensed Optician")
add_paragraph("📍 Jeddah, Saudi Arabia    📞 +966 540 919 725    ✉️ hasin3112@gmail.com")

# Objective
add_section_title("Objective")
add_paragraph("To leverage my extensive experience of over 14 years as a licensed Optician and Optical Store Manager, combining leadership, customer service excellence, and technical expertise to contribute effectively to the success and growth of a leading optical organization.")

# Skills
add_section_title("Key Skills")
add_paragraph("Technical Skills", bold=True)
for skill in [
    "Expert in Optical Lenses & Products",
    "Contact Lenses Fitting & Sales",
    "Medical Insurance Handling",
    "Eyewear Adjustments & Repairs",
    "Visual Aids & Refraction Assistance"
]:
    add_bullet_point(skill)

add_paragraph("Professional Skills", bold=True)
for skill in [
    "Customer Service Excellence",
    "Sales Target Management",
    "Team Leadership & Training",
    "Problem-Solving & Decision-Making",
    "Strategic Workflow Management",
    "Attention to Detail"
]:
    add_bullet_point(skill)

# Experience
add_section_title("Professional Experience")
add_paragraph("Optician & Branch Manager", bold=True)
add_paragraph("Hussam Optics Company – Jeddah, Saudi Arabia | 2011 – Present")
for duty in [
    "Began as an Optician and Sales Manager.",
    "Promoted to Branch Manager overseeing daily operations.",
    "Responsible for team management, customer service, inventory control, and achieving sales targets.",
    "Implemented workflow strategies that improved customer satisfaction and branch efficiency."
]:
    add_bullet_point(duty)

add_paragraph("Optician & Sales Assistant", bold=True)
add_paragraph("El-Nour Optics – Egypt | 2008 – 2011")
for duty in [
    "Assisted customers with selecting eyewear and lenses.",
    "Supported optical testing and sales operations.",
    "Contributed to increasing sales through effective communication and product knowledge."
]:
    add_bullet_point(duty)

# Education
add_section_title("Education & Certifications")
education_items = [
    "Diploma in Ophthalmic Medical Assisting – Magrabi Foundation – Sep 2023 to May 2024",
    "Digital Marketing Challenger Program – Udacity – Feb 2023 to Mar 2023",
    "Data Analysis Program – Udacity – Aug 2022 to Sep 2022",
    "Optician Diploma – Optical Institute – Sep 2006 to Sep 2008"
]
for item in education_items:
    add_bullet_point(item)

# Licenses
add_section_title("Licenses")
licenses = [
    "Licensed Optician – Saudi Commission for Health Specialties – Since 2011",
    "Licensed Optician – Egyptian Ministry of Health and Population – Since 2009"
]
for lic in licenses:
    add_bullet_point(lic)

# Communication
add_section_title("Communication Skills")
for point in [
    "Excellent written and verbal communication",
    "Ability to collaborate across teams and departments",
    "Provide clear, constructive feedback and guidance"
]:
    add_bullet_point(point)

# Leadership
add_section_title("Leadership & Management")
for point in [
    "Led and developed high-performing optical teams",
    "Implemented strategies to improve store performance",
    "Skilled in decision-making, mentoring, and motivating staff"
]:
    add_bullet_point(point)

# Notes
add_section_title("Notes")
add_paragraph("✔ Available upon request: references, license documents, and training certificates.")
add_paragraph("✔ Open to relocation and flexible work hours.")

# Save document
file_path = "/mnt/data/ELHUSSEINI_Optician_CV.docx"
doc.save(file_path)

file_path
